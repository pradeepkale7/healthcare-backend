from fastapi import APIRouter, UploadFile, File, HTTPException, Depends
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.models.file_import import FileImport
from app.services.storage_service import save_file
from datetime import datetime
from app.services.llm_service import generate_mapping_with_llm

import pandas as pd
from io import BytesIO
from typing import List, Dict, Any
import numpy as np
import os

# PDF processing imports
import PyPDF2
import pdfplumber
from pdfminer.high_level import extract_text as pdfminer_extract_text

# DOCX processing imports
import docx
from docx import Document
import xml.etree.ElementTree as ET

# Additional imports for table extraction
import tabula
import camelot

router = APIRouter()


def extract_form_data_from_text(text: str) -> Dict[str, Any]:
    """
    Extract form-like data from text content
    Looks for key-value pairs, labels, and form fields
    """
    form_data = {
        "headers": [],
        "sample_data": [],
        "form_fields": {}
    }
    
    import re
    
    # Common form patterns
    patterns = [
        r'([A-Za-z\s]+):\s*([^\n\r]+)',  # Label: Value
        r'([A-Za-z\s]+)\s*=\s*([^\n\r]+)',  # Label = Value
        r'([A-Za-z\s]+)\s*-\s*([^\n\r]+)',  # Label - Value
        r'([A-Za-z\s]+)\s*\|\s*([^\n\r]+)',  # Label | Value
        r'([A-Za-z\s]+)\s*\.\s*([^\n\r]+)',  # Label . Value
    ]
    
    found_fields = {}
    
    for pattern in patterns:
        matches = re.findall(pattern, text, re.MULTILINE)
        for match in matches:
            key = match[0].strip()
            value = match[1].strip()
            
            # Filter out likely non-form content
            if (len(key) > 2 and len(key) < 50 and 
                len(value) > 0 and len(value) < 200 and
                not key.lower().startswith(('page', 'figure', 'table', 'section'))):
                found_fields[key] = value
    
    if found_fields:
        form_data["headers"] = list(found_fields.keys())
        form_data["sample_data"] = [found_fields]
        form_data["form_fields"] = found_fields
    
    return form_data


def extract_pdf_content(file_contents: bytes) -> Dict[str, Any]:
    """
    Extract content from PDF files with priority order:
    1. Form data extraction
    2. General text content
    3. Table extraction (last resort)
    """
    pdf_data = {
        "text_content": "",
        "tables": [],
        "headers": [],
        "sample_data": [],
        "form_fields": {},
        "extraction_method": ""
    }
    
    try:
        # Step 1: Extract all text content first
        with pdfplumber.open(BytesIO(file_contents)) as pdf:
            full_text = ""
            tables = []
            
            for page in pdf.pages:
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    full_text += page_text + "\n"
                
                # Also collect tables for later use
                page_tables = page.extract_tables()
                if page_tables:
                    tables.extend(page_tables)
            
            pdf_data["text_content"] = full_text
            pdf_data["tables"] = tables
            
            # Step 2: Try to extract form data first
            if full_text.strip():
                form_data = extract_form_data_from_text(full_text)
                if form_data["headers"]:
                    pdf_data["headers"] = form_data["headers"]
                    pdf_data["sample_data"] = form_data["sample_data"]
                    pdf_data["form_fields"] = form_data["form_fields"]
                    pdf_data["extraction_method"] = "form_extraction"
                    print(f"Form data extracted: {len(form_data['headers'])} fields found")
                    return pdf_data
            
            # Step 3: If no form data, try general text processing
            if full_text.strip():
                # Look for structured text patterns (not form-like)
                lines = [line.strip() for line in full_text.split('\n') if line.strip()]
                
                # Try to identify if text has structured information
                structured_content = []
                for line in lines:
                    if len(line) > 10 and any(char in line for char in [',', '|', '\t']):
                        structured_content.append(line)
                
                if structured_content:
                    # Process as structured text
                    pdf_data["headers"] = ["structured_content"]
                    pdf_data["sample_data"] = [{"structured_content": '\n'.join(structured_content[:5])}]
                    pdf_data["extraction_method"] = "structured_text"
                    print("Structured text content extracted")
                    return pdf_data
            
            # Step 4: Finally, try table extraction as last resort
            if tables:
                first_table = tables[0]
                if len(first_table) > 0:
                    headers = [str(cell).strip() if cell else f"Column_{i}" 
                              for i, cell in enumerate(first_table[0])]
                    pdf_data["headers"] = headers
                    
                    # Get sample rows (skip header row)
                    sample_rows = []
                    for row in first_table[1:6]:  # Get up to 5 sample rows
                        row_dict = {}
                        for i, cell in enumerate(row):
                            if i < len(headers):
                                row_dict[headers[i]] = str(cell).strip() if cell else None
                        sample_rows.append(row_dict)
                    pdf_data["sample_data"] = sample_rows
                    pdf_data["extraction_method"] = "table_extraction"
                    print("Table data extracted as last resort")
                    return pdf_data
            
            # If nothing structured found, return raw text
            if full_text.strip():
                pdf_data["headers"] = ["extracted_text"]
                pdf_data["sample_data"] = [{"extracted_text": full_text[:1000]}]
                pdf_data["extraction_method"] = "raw_text"
                print("Raw text extracted")
                return pdf_data
            
    except Exception as e:
        print(f"pdfplumber failed: {e}")
        
        # Fallback methods
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_contents))
            text_content = ""
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
            
            if text_content.strip():
                # Try form extraction on fallback text
                form_data = extract_form_data_from_text(text_content)
                if form_data["headers"]:
                    pdf_data["headers"] = form_data["headers"]
                    pdf_data["sample_data"] = form_data["sample_data"]
                    pdf_data["form_fields"] = form_data["form_fields"]
                    pdf_data["extraction_method"] = "form_extraction_fallback"
                else:
                    pdf_data["headers"] = ["extracted_text"]
                    pdf_data["sample_data"] = [{"extracted_text": text_content[:1000]}]
                    pdf_data["extraction_method"] = "raw_text_fallback"
                
                pdf_data["text_content"] = text_content
                
        except Exception as e2:
            print(f"PyPDF2 failed: {e2}")
            
            try:
                text_content = pdfminer_extract_text(BytesIO(file_contents))
                pdf_data["text_content"] = text_content
                pdf_data["headers"] = ["extracted_text"]
                pdf_data["sample_data"] = [{"extracted_text": text_content[:1000]}]
                pdf_data["extraction_method"] = "raw_text_pdfminer"
            except Exception as e3:
                print(f"All PDF extraction methods failed: {e3}")
                raise HTTPException(status_code=500, detail="Failed to extract content from PDF")
    
    return pdf_data


def extract_docx_content(file_contents: bytes) -> Dict[str, Any]:
    """
    Extract content from DOCX files with priority order:
    1. Form data extraction
    2. General text content
    3. Table extraction (last resort)
    """
    docx_data = {
        "text_content": "",
        "tables": [],
        "headers": [],
        "sample_data": [],
        "form_fields": {},
        "extraction_method": ""
    }
    
    try:
        doc = Document(BytesIO(file_contents))
        
        # Extract all text content first
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        text_content = "\n".join(full_text)
        docx_data["text_content"] = text_content
        
        # Step 1: Try to extract form data first
        if text_content.strip():
            form_data = extract_form_data_from_text(text_content)
            if form_data["headers"]:
                docx_data["headers"] = form_data["headers"]
                docx_data["sample_data"] = form_data["sample_data"]
                docx_data["form_fields"] = form_data["form_fields"]
                docx_data["extraction_method"] = "form_extraction"
                print(f"Form data extracted from DOCX: {len(form_data['headers'])} fields found")
                return docx_data
        
        # Step 2: Try structured text processing
        if text_content.strip():
            lines = [line.strip() for line in text_content.split('\n') if line.strip()]
            
            # Look for structured content patterns
            structured_content = []
            for line in lines:
                if len(line) > 10 and any(char in line for char in [',', '|', '\t', ':']):
                    structured_content.append(line)
            
            if structured_content:
                docx_data["headers"] = ["structured_content"]
                docx_data["sample_data"] = [{"structured_content": '\n'.join(structured_content[:5])}]
                docx_data["extraction_method"] = "structured_text"
                print("Structured text content extracted from DOCX")
                return docx_data
        
        # Step 3: Extract tables as last resort
        if doc.tables:
            tables_data = []
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            docx_data["tables"] = tables_data
            
            # Use first table for headers and sample data
            if tables_data and len(tables_data[0]) > 0:
                first_table = tables_data[0]
                headers = [str(cell).strip() if cell else f"Column_{i}" 
                          for i, cell in enumerate(first_table[0])]
                docx_data["headers"] = headers
                
                # Get sample rows
                sample_rows = []
                for row in first_table[1:6]:  # Skip header, get up to 5 rows
                    row_dict = {}
                    for i, cell in enumerate(row):
                        if i < len(headers):
                            row_dict[headers[i]] = str(cell).strip() if cell else None
                    sample_rows.append(row_dict)
                docx_data["sample_data"] = sample_rows
                docx_data["extraction_method"] = "table_extraction"
                print("Table data extracted from DOCX as last resort")
                return docx_data
        
        # Step 4: If nothing structured found, return raw text
        if text_content.strip():
            docx_data["headers"] = ["extracted_text"]
            docx_data["sample_data"] = [{"extracted_text": text_content[:1000]}]
            docx_data["extraction_method"] = "raw_text"
            print("Raw text extracted from DOCX")
            return docx_data
    
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        raise HTTPException(status_code=500, detail="Failed to extract content from DOCX file")
    
    return docx_data


@router.post("/upload")
async def upload_file(file: UploadFile = File(...), db: Session = Depends(get_db)):
    if not file.filename.endswith((".csv", ".xlsx", ".tsv", ".pdf", ".docx")):
        raise HTTPException(status_code=400, detail="Unsupported file type")
    
    PREDEFINED_COLUMNS = [
        "member_id", "first_name", "last_name", "dob", "gender", "email", "phone", "address",
        "npi_number", "provider_name", "policy_number", "plan_name", "group_number",
        "policy_start_date", "policy_end_date", "claim_date", "admission_date", "discharge_date",
        "amount_claimed", "amount_approved", "claim_status", "rejection_reason", "diagnosis_code", "diagnosis_description"
    ]

    file_contents = await file.read()
    
    storage_type, result = save_file(file, file.filename)

    if storage_type == "local":
        local_path = result
        s3_bucket = s3_key = None
    else:
        local_path = None
        s3_bucket, s3_key = result

    new_import = FileImport(
        filename=file.filename,
        file_extension=file.filename.split(".")[-1],
        storage_type=storage_type,
        local_path=local_path,
        s3_bucket=s3_bucket,
        s3_key=s3_key,
        processing_status="Uploaded",
        upload_time=datetime.utcnow()
    )

    db.add(new_import)
    db.commit()
    
    # Get file extension
    extension = file.filename.split(".")[-1].lower()
    
    headers = []
    sample_rows = []
    extracted_content = ""

    try:
        if extension == "csv":
            df = pd.read_csv(BytesIO(file_contents), delimiter=",")
            df = df.replace({np.nan: None})
            headers = list(df.columns)
            sample_rows = df.to_dict(orient="records")
            
        elif extension == "tsv":
            df = pd.read_csv(BytesIO(file_contents), delimiter="\t")
            df = df.replace({np.nan: None})
            headers = list(df.columns)
            sample_rows = df.to_dict(orient="records")
            
        elif extension in ["xls", "xlsx"]:
            try:
                df = pd.read_excel(BytesIO(file_contents), engine="openpyxl")
            except:
                df = pd.read_excel(BytesIO(file_contents), engine="xlrd")
            
            df = df.replace({np.nan: None})
            headers = list(df.columns)
            sample_rows = df.to_dict(orient="records")
            
        elif extension == "pdf":
            pdf_data = extract_pdf_content(file_contents)
            headers = pdf_data["headers"]
            sample_rows = pdf_data["sample_data"]
            extracted_content = pdf_data["text_content"]
            extraction_method = pdf_data.get("extraction_method", "unknown")
            
            print(f"PDF processed using method: {extraction_method}")
            
            # If form fields were extracted, include them in response
            form_fields = pdf_data.get("form_fields", {})
            
        elif extension == "docx":
            docx_data = extract_docx_content(file_contents)
            headers = docx_data["headers"]
            sample_rows = docx_data["sample_data"]
            extracted_content = docx_data["text_content"]
            extraction_method = docx_data.get("extraction_method", "unknown")
            
            print(f"DOCX processed using method: {extraction_method}")
            
            # If form fields were extracted, include them in response
            form_fields = docx_data.get("form_fields", {})
                
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        
        # Validation for structured data (CSV, TSV, XLSX)
        if extension in ["csv", "tsv", "xlsx", "xls"]:
            if len(pd.DataFrame(sample_rows)) == 0:
                raise HTTPException(status_code=400, detail="File is empty or has no valid data")
            if len(headers) == 0:
                raise HTTPException(status_code=400, detail="File has no headers")
            if len(headers) > 50:
                raise HTTPException(status_code=400, detail="File has too many headers (max 50 allowed)")
            if len(headers) != len(set(headers)):
                raise HTTPException(status_code=400, detail="File has duplicate headers")
        
        # Generate mapping using LLM service
        result = generate_mapping_with_llm(headers, sample_rows)
        print(f"Generated mapping for {file.filename}: {result}")
        
        response_data = {
            "file_import_id": new_import.import_id,
            "filename": file.filename,
            "headers": headers,
            "predefined_columns": PREDEFINED_COLUMNS,
            "sample_data": sample_rows[:10],  # Limit sample data to first 10 rows
            "mapping_result": result
        }
        
        # Add extracted content and processing info for PDF/DOCX files
        if extension in ["pdf", "docx"]:
            response_data["extracted_content"] = extracted_content[:2000]  # First 2000 characters
            response_data["extraction_method"] = extraction_method
            
            # Include form fields if they were extracted
            if 'form_fields' in locals() and form_fields:
                response_data["form_fields"] = form_fields
                response_data["total_form_fields"] = len(form_fields)
            
        return response_data

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to process file: {str(e)}")